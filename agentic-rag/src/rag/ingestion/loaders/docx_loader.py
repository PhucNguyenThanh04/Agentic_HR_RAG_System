from io import BytesIO
from typing import Any

from docx import Document as DocxDocument
from docx.oxml.ns import qn

from src.rag.ingestion.loaders.base_loader import BaseLoader, Document, UploadFile
from src.core.setup_logging import setup_logger

logger = setup_logger(__name__)


def _get_heading_level(paragraph) -> int | None:
    """Trả về heading level (1-9) hoặc None nếu không phải heading."""
    style_name = paragraph.style.name
    if style_name.startswith("Heading "):
        try:
            return int(style_name.split(" ")[1])
        except (IndexError, ValueError):
            pass
    # DOCX đôi khi dùng outline level thay vì style name
    outline = paragraph._element.find(qn("w:outlineLvl"))
    if outline is not None:
        val = outline.get(qn("w:val"))
        if val is not None:
            return int(val) + 1
    return None


class DocxLoader(BaseLoader):
    """
    Load DOCX, group nội dung theo heading structure.

    Mỗi section (từ Heading đến Heading tiếp theo) = 1 Document.
    Phù hợp với nội quy công ty có structure rõ ràng:
        Heading 1: "Chương I: Quy định chung"
        Heading 2: "Điều 1: Phạm vi áp dụng"
        Body:       "..."
    """

    def load(
        self,
        file: UploadFile,
        allowed_roles: list[Any] | None = None,
        extra_metadata: dict[str, Any] | None = None,
    ) -> list[Document]:
        filename = self._upload_filename(file, extra_metadata)

        try:
            self._reset_upload_file(file)
            content = file.file.read()
            docx = DocxDocument(BytesIO(content))
        except Exception as e:
            raise ValueError(f"Cannot read DOCX '{filename}': {e}") from e

        docs: list[Document] = []

        current_heading: str = ""           # heading của section hiện tại
        current_heading_level: int = 0
        current_paragraphs: list[str] = []  # body text tích lũy
        section_index: int = 0

        def _flush():
            """Đẩy section hiện tại vào docs list."""
            nonlocal section_index
            text = "\n".join(current_paragraphs).strip()
            if self._is_empty(text):
                return
            docs.append(Document(
                content=text,
                metadata=self._build_metadata(
                    {
                        "source_file":     filename,
                        "doc_type":        "docx",
                        "section":         current_heading,
                        "heading_level":   current_heading_level,
                        "section_index":   section_index,
                    },
                    allowed_roles,
                    extra_metadata,
                ),
            ))
            section_index += 1

        for para in docx.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            level = _get_heading_level(para)

            if level is not None:
                # Gặp heading mới → flush section cũ, bắt đầu section mới
                _flush()
                current_paragraphs = []
                current_heading = text
                current_heading_level = level
            else:
                current_paragraphs.append(text)

        # Flush section cuối cùng
        _flush()

        if not docs:
            logger.warning(f"'{filename}' has no extractable content.")
            raise ValueError(f"'{filename}' has no extractable content.")

        logger.info(f"Loaded {len(docs)} sections from DOCX '{filename}'")
        return docs
